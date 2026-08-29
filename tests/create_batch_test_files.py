import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz

os.makedirs('verification_evidence/batch_input', exist_ok=True)
os.makedirs('verification_evidence/batch_output', exist_ok=True)

# 1. Create DOCX with Table, Custom Fonts, and Colors
doc = docx.Document()

title = doc.add_heading('CONFIDENTIAL VENDOR AGREEMENT', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('This document contains formatted table data and proprietary company records.')
p.runs[0].font.size = Pt(11)

# Add a 3x3 styled table
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Department'
hdr_cells[1].text = 'Contact Email (PII)'
hdr_cells[2].text = 'Emergency Phone (PII)'

row1_cells = table.rows[1].cells
row1_cells[0].text = 'Engineering'
row1_cells[1].text = 'eng-lead@techcorp-innovations.com'
row1_cells[2].text = '555-888-1234'

row2_cells = table.rows[2].cells
row2_cells[0].text = 'Executive'
row2_cells[1].text = 'ceo-private@techcorp-innovations.com'
row2_cells[2].text = '555-999-5678'

docx_path = "verification_evidence/batch_input/sample_table.docx"
doc.save(docx_path)
print(f"[1] Created DOCX with table: {docx_path}")

# 2. Create Two-Column PDF
pdf = fitz.open()
page = pdf.new_page(width=612, height=792)

page.insert_text((50, 60), "TWO-COLUMN EXECUTIVE BRIEFING", fontsize=14, fontname="helv")

# Column 1 (Left: x=50 to 280)
col1_lines = [
    "COLUMN 1: REGIONAL OPERATIONS",
    "Headquarters: New York Office",
    "Director: Johnathan Davis",
    "Desk Phone: +1-212-555-0188",
    "Email: ny-desk@enterprise-corp.net",
    "Status: Audited and verified."
]
y = 100
for l in col1_lines:
    page.insert_text((50, y), l, fontsize=10, fontname="helv")
    y += 24

# Column 2 (Right: x=330 to 560)
col2_lines = [
    "COLUMN 2: GLOBAL COMPLIANCE",
    "Branch: London Operations",
    "Officer: Victoria Sterling",
    "Desk Phone: +44-20-7946-0199",
    "Email: uk-compliance@enterprise-corp.net",
    "Status: Certified compliant."
]
y = 100
for l in col2_lines:
    page.insert_text((330, y), l, fontsize=10, fontname="helv")
    y += 24

# Draw vertical column dividing line
page.draw_line((305, 90), (305, 260), color=(0.7, 0.7, 0.7), width=1)

pdf_path = "verification_evidence/batch_input/two_column.pdf"
pdf.save(pdf_path)
pdf.close()
print(f"[2] Created 2-column PDF: {pdf_path}")

# 3. Create Plain Text file
txt_path = "verification_evidence/batch_input/plain_notes.txt"
with open(txt_path, "w", encoding="utf-8") as f:
    f.write("CONFIDENTIAL MEMO\nAuthorized Contact: confidential-officer@agency.gov\nDirect: 555-123-9876\n")
print(f"[3] Created TXT file: {txt_path}")
