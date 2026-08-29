import os
import docx
from playwright.sync_api import sync_playwright

def docx_to_html(docx_path):
    doc = docx.Document(docx_path)
    html = "<div style='font-family: Arial, sans-serif; padding: 20px; background: white; color: black; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);'>"
    for p in doc.paragraphs:
        if p.text.strip():
            align = p.alignment or 0
            align_str = "center" if align == 1 else "left"
            html += f"<p style='text-align: {align_str}; font-weight: {'bold' if 'AGREEMENT' in p.text else 'normal'};'>{p.text}</p>"
    for table in doc.tables:
        html += "<table style='border-collapse: collapse; width: 100%; margin-top: 15px; border: 1px solid #ccc;'>"
        for r_idx, row in enumerate(table.rows):
            html += "<tr>"
            for cell in row.cells:
                is_hdr = (r_idx == 0)
                bg = "#f3f4f6" if is_hdr else "white"
                fw = "bold" if is_hdr else "normal"
                html += f"<td style='border: 1px solid #ccc; padding: 8px 12px; background: {bg}; font-weight: {fw}; font-size: 13px;'>{cell.text}</td>"
            html += "</tr>"
        html += "</table>"
    html += "</div>"
    return html

orig_html = docx_to_html("verification_evidence/batch_input/sample_table.docx")
redacted_html = docx_to_html("verification_evidence/batch_output/extracted/sample_table_redacted.docx")

comparison_html = f"""
<!DOCTYPE html>
<html>
<head>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #0f172a; color: white; padding: 30px; }}
h2 {{ margin-bottom: 20px; text-align: center; }}
.container {{ display: flex; gap: 24px; justify-content: center; }}
.col {{ flex: 1; max-width: 580px; }}
.label {{ font-size: 14px; font-weight: bold; margin-bottom: 8px; color: #94a3b8; text-transform: uppercase; letter-spacing: 0.05em; }}
</style>
</head>
<body>
<h2>DOCX Batch Export Table Preservation: Side-by-Side Comparison</h2>
<div class="container">
  <div class="col">
    <div class="label">Original DOCX (Before Redaction)</div>
    {orig_html}
  </div>
  <div class="col">
    <div class="label">Exported DOCX (After Format-Preserved Redaction)</div>
    {redacted_html}
  </div>
</div>
</body>
</html>
"""

html_path = "verification_evidence/docx_comparison.html"
with open(html_path, "w", encoding="utf-8") as f:
    f.write(comparison_html)

chrome_path = r"C:\Program Files\Google\Chrome\Application\chrome.exe"
with sync_playwright() as p:
    browser = p.chromium.launch(executable_path=chrome_path, headless=True)
    page = browser.new_page(viewport={'width': 1280, 'height': 600})
    page.goto(f"file:///{os.path.abspath(html_path)}")
    page.screenshot(path="verification_evidence/batch_docx_comparison.png")
    browser.close()

print("Saved verification_evidence/batch_docx_comparison.png")
